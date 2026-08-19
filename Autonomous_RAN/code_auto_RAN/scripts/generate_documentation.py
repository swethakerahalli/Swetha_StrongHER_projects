#!/usr/bin/env python3
"""Generate DOCX and PDF from AUTONOMOUS_RAN_IMPLEMENTATION.md."""

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

DOCS = ROOT / "docs"
MD_PATH = DOCS / "AUTONOMOUS_RAN_IMPLEMENTATION.md"
DOCX_PATH = DOCS / "AUTONOMOUS_RAN_IMPLEMENTATION.docx"
PDF_PATH = DOCS / "AUTONOMOUS_RAN_IMPLEMENTATION.pdf"


def _read_md() -> str:
    return MD_PATH.read_text(encoding="utf-8")


def generate_docx(text: str) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("|") and "---" not in stripped:
            doc.add_paragraph(stripped)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("```"):
            continue
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            doc.add_paragraph(clean)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def generate_pdf(text: str) -> Path:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=9)

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("```"):
            continue
        try:
            if stripped.startswith("# "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.multi_cell(0, 7, stripped[2:][:200])
                pdf.ln(2)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.multi_cell(0, 6, stripped[3:][:200])
                pdf.ln(1)
            elif stripped.startswith("### "):
                pdf.set_font("Helvetica", "B", 10)
                pdf.multi_cell(0, 5, stripped[4:][:200])
            else:
                pdf.set_font("Helvetica", size=8)
                clean = stripped.encode("ascii", errors="ignore").decode("ascii")
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
                clean = clean.replace("|", " ").replace("─", "-").replace("→", "->")[:250]
                if clean:
                    pdf.multi_cell(0, 4, clean)
        except Exception:
            pdf.add_page()
            continue

    pdf.output(PDF_PATH)
    return PDF_PATH


def main():
    if not MD_PATH.exists():
        print(f"Missing {MD_PATH}")
        sys.exit(1)
    text = _read_md()
    DOCS.mkdir(parents=True, exist_ok=True)
    print("Generating DOCX...")
    generate_docx(text)
    print(f"  -> {DOCX_PATH}")
    print("Generating PDF...")
    generate_pdf(text)
    print(f"  -> {PDF_PATH}")
    print("Done.")


if __name__ == "__main__":
    main()
