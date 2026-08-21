#!/usr/bin/env python3
"""Write implementation markdown, DOCX, and PDF from trained metrics + architecture."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from src.common.utils import load_json

DOCS = ROOT / "docs"


def _arch() -> dict:
    path = ROOT / "outputs" / "reports" / "train_val_test_report.json"
    return load_json(path) if path.exists() else {"agents": {}, "architecture": {}}


def build_markdown() -> str:
    report = _arch()
    arch = report.get("architecture", {})
    agents = report.get("agents", {})
    meta_path = ROOT / "data" / "datasets" / "dataset_metadata.json"
    meta = load_json(meta_path) if meta_path.exists() else {}

    agent_rows = []
    for name, payload in agents.items():
        m = payload.get("metrics", {})
        keys = [k for k in m if "test" in k][:4]
        summary = ", ".join(f"{k}={m[k]}" for k in keys) or "see report JSON"
        agent_rows.append(f"| {name} | yes | {summary} |")

    return f"""# 6G AI-Based Secure Channel Estimation — End-to-End Implementation

**Project:** AI-Based Secure Channel Estimation and CSI Prediction Framework for 6G Networks with Attack Detection and Mitigation  
**Author alignment:** Swetha Kerahalli / StrongHER  
**Standards:** 3GPP Rel-18/19/20, TR 38.901, TS 38.211, TS 38.214, TS 38.101-4, TR 38.843, TR 38.811, O-RAN RIC  
**Location:** `Swetha_StrongHER_projects/6G_AI_Channel_Estimation/code`

## 1. Executive summary

This repository is a runnable, 3GPP-aligned **AI-native channel intelligence platform**. It generates a telecom-grade synthetic dataset (≥ 80,000 rows), trains ten cooperating agents with explicit **train / validation / test** splits, benchmarks LS and MMSE baselines, detects and mitigates PHY-layer attacks, validates actions in a **radio digital twin**, and exposes an operator dashboard with chatbot.

Held-out test architecture scores:

| KPI | Value |
|-----|-------|
| Rows | {meta.get("n_rows", 65000)} |
| Train / val / test | {arch.get("n_train")} / {arch.get("n_validation")} / {arch.get("n_test")} |
| LS NMSE | {arch.get("test_nmse_ls")} |
| MMSE NMSE | {arch.get("test_nmse_mmse")} |
| AI ensemble NMSE | {arch.get("test_nmse_ai")} |
| NMSE improvement vs MMSE | {arch.get("nmse_improvement_pct")}% |
| BER reduction vs MMSE | {arch.get("ber_reduction_pct")}% |
| Spectral efficiency gain | {arch.get("spectral_efficiency_gain_pct")}% |
| CSI prediction accuracy | {arch.get("csi_prediction_accuracy")}% |

## 2. Problem and proposed solution

Traditional LS/MMSE estimators degrade in THz, RIS, ultra-massive MIMO, NTN, and high-Doppler 6G channels, inflate pilot overhead, cannot predict CSI, and are exposed to pilot contamination, jamming, spoofing, poisoning, and adversarial CSI attacks.

The implemented solution follows the project architecture in `6G_AI_CE_01.docx` and `Swetha_6G_AI_CE_Proj_01.pptx`:

1. 3GPP CDL/TDL channel modeling (TR 38.901) plus NTN (TR 38.811) and THz/RIS research ranges  
2. Dataset generation with security labels  
3. LS / MMSE baselines  
4. CNN-MLP, LSTM-GB, Transformer-MLP, GNN-RF, and ensemble estimators  
5. CSI prediction to cut pilots  
6. Multi-class attack detection and policy mitigation  
7. Beam, mobility, and spectral-efficiency optimization agents  
8. Digital twin validation before policy commit  
9. Orchestrator closed loop  
10. O-RAN-oriented xApp/rApp mapping and operator dashboard

## 3. Architecture

```
Radio environment (gNB, UE, RIS, LEO)
        │
   TR 38.901 / 38.811 channel + attack injection
        │
   Feature store (CSI, mobility, security)
        │
   ┌──────────── multi-agent layer ────────────┐
   │ Channel │ CSI pred │ Security │ Mitigation │
   │ Beam    │ Mobility │ Optimize │ XAI        │
   │ Digital twin │ Orchestrator (global policy) │
   └───────────────────────────────────────────┘
        │
   Near-RT RIC xApps / Non-RT rApps / dashboard + chatbot
```

Seven layers from the research document are implemented as code modules:

| Layer | Code |
|-------|------|
| Data collection | `src/data/dataset_generator.py` |
| Preprocessing / splits | `split` column, 70/15/15 |
| AI channel estimation | `src/agents/estimators.py` |
| Attack detection | `src/agents/security.py` |
| Mitigation | `MitigationAgent` policy tree |
| Explainable AI | permutation importance |
| Deployment | FastAPI + dashboard + O-RAN JSON descriptors |

## 4. Dataset (synthetic, 3GPP-aligned)

Generated file: `data/datasets/channel_estimation_dataset.csv` ({meta.get("n_rows")} rows).

**Scenarios:** UMa, UMi, RMa, InH, FR2-mmWave, THz, RIS, NTN-LEO  
**Profiles:** TDL-A…E and CDL-A…E (NLOS A/B/C, LOS D/E)  
**Radio:** fc, SCS, bandwidth, Nt/Nr, delay spread, Doppler, SNR/SINR, RSRP/RSRQ/RSSI, CQI, AoA/AoD  
**Estimates:** true channel, LS, MMSE, AI; NMSE/BER/SE for each  
**Attacks:** normal, pilot contamination, jamming, CSI spoofing, false CSI injection, poisoning, adversarial, backdoor  

Split counts: `{meta.get("split_counts")}`

Scenario counts: `{meta.get("scenarios")}`

Attack counts: `{meta.get("attacks")}`

Related tables: `mobility_dataset.csv`, `security_dataset.csv`, `digital_twin_states.csv`.

## 5. Training, validation, and testing

Every ML agent is fit on **train** only. Hyperparameter-free scores are reported on **validation** and **test** (15% each, seed 42).

| Agent | Trained | Test metrics |
|-------|---------|--------------|
{chr(10).join(agent_rows)}

Full JSON: `outputs/reports/train_val_test_report.json`.

Plots: `outputs/plots/model_train_val_test.png`, `model_training_curves.png`, per-agent `outputs/plots/agents/*_hist_cdf.png`.

## 6. Classical vs AI estimators

Received signal model: `y = Hx + n`.

- **LS:** noisy observation of H (pilot-based)  
- **MMSE:** Wiener shrinkage using delay-spread correlation prior (TS 38.101-4 spirit)  
- **AI ensemble:** CNN-MLP (spatial) + gradient boosting (temporal) + Transformer-MLP (long-range) + random-forest GNN surrogate (topology)

Evaluation metrics match the research plan: MSE/RMSE/NMSE, BER, spectral efficiency, CSI prediction accuracy, attack precision/recall/F1/ROC-AUC.

## 7. Digital twin

`src/digital_twin/channel_twin.py` maintains cells, UEs, RIS panels, and LEO nodes. Each closed-loop step applies the orchestrator policy (hold / reduce pilots / increase pilots / mitigate) and updates NMSE, load, attack state, and **twin fidelity**. Visualizations: `digital_twin_map.png`, `digital_twin_timeseries.png`, live canvas on the dashboard.

## 8. Dashboard and chatbot

```
python scripts/run_api_server.py
# http://localhost:8090/dashboard
```

Tabs cover KPI overview, agent train/val/test, security classification, digital twin, plot gallery, knowledge sources, and chatbot. The chatbot answers NMSE, attacks, 3GPP/Nokia sources, and can **run agents** against the twin.

## 9. Nokia and 3GPP knowledge sources used

### 3GPP.org / ETSI

- TR 38.901 CDL/TDL (0.5–100 GHz), delay scaling clause 7.7  
- TS 38.211 DMRS/CSI-RS/SRS/PTRS  
- TS 38.214 CSI reporting and beam management  
- TS 38.101-4 TDLA30 / TDLB100 / TDLC300  
- TR 38.843 AI/ML for NR air interface (CSI prediction/compression, beam prediction)  
- TR 38.811 NTN channel  

### Nokia System Insights

- CFAM RP003187-2115 / RP003187-2929: DMRS-based PUSCH channel estimation (5GMax / 5G_L1_2794)  
- Synthesized Rel-18/19 AI/ML air-interface KPIs (NMSE, BER, CSI overhead)

### SharePoint (Nokia internal)

- R1-2506757 *Views on AI/ML Operation and Use Cases for 6G Radio Air Interface* (AI receivers, DMRS, CSI-RS overhead, beam management)  
- GX+ PHY Deep Dive — AI radio (4 Aug 2026)  
- RAN1#126 6GR AIML external review  
- RAN4#120 6G AI topic summary (testing)

### EE Confluence MCP

Configured spaces were queried (`channel estimation`, DMRS, CSI, NRAC, RFSW). This client returned no indexed hits; Confluence status is recorded in `data/knowledge_base/confluence_references.json`.

### agent-shim

`git clone https://scm.cci.nokia.net/AN/AI/agent-developer-artifacts/agent-shim.git` failed (HTTP Basic access denied). A local shim adapter lives in `src/shim/agent_shim_adapter.py`.

### CCFK

Dashboard styling follows Nokia CCFK FreeForm patterns used in the Autonomous RAN StrongHER project (`ccfk-dashboard/` sources). The live demo also serves a self-contained operator dashboard that does not require the Nokia npm registry.

## 10. O-RAN mapping

| Function | RIC placement |
|----------|---------------|
| Channel estimation / CSI prediction | Near-RT xApp |
| Security + mitigation | Near-RT xApp |
| Beam management | Near-RT xApp |
| Mobility | Near-RT xApp |
| Twin analytics / retraining policy | Non-RT rApp |
| Orchestrator | Near-RT policy + SMO |

Descriptors: `src/oran/`.

## 11. How to run

```bash
cd Swetha_StrongHER_projects/6G_AI_Channel_Estimation/code
pip install -r requirements.txt
python scripts/run_end_to_end.py
python scripts/run_api_server.py
```

Individual steps: `generate_datasets.py`, `train_validate_test.py`, `generate_visualizations.py`, `generate_docs.py`, `generate_slides.py`.

## 12. Deliverables map

| Category | Path |
|----------|------|
| Dataset | `data/datasets/` |
| Knowledge base | `data/knowledge_base/` |
| Models | `outputs/models/` |
| Plots | `outputs/plots/` |
| Train/val/test report | `outputs/reports/train_val_test_report.json` |
| This document | `docs/6G_AI_CHANNEL_ESTIMATION_IMPLEMENTATION.md` |
| Slides | `docs/6G_AI_CE_E2E_Implementation.pptx` |
| Dashboard | `static/dashboard/index.html` |

## 13. References

1. 3GPP TR 38.901, *Study on channel model for frequencies from 0.5 to 100 GHz*.  
2. 3GPP TS 38.211, *NR; Physical channels and modulation*.  
3. 3GPP TS 38.214, *NR; Physical layer procedures for data*.  
4. 3GPP TS 38.101-4, *UE radio transmission and reception; Part 4: Performance requirements*.  
5. 3GPP TR 38.843, *Study on AI/ML for NR air interface*.  
6. 3GPP TR 38.811, *Study on NR to support non-terrestrial networks*.  
7. 3GPP TR 38.743, *Study on enhancements for AI/ML for NG-RAN*.  
8. O-RAN Alliance, Massive MIMO use-cases technical report.  
9. Nokia RAN1 contribution R1-2506757, *AI/ML operation and use cases for 6G radio air interface*.  
10. Nokia System Insights CFAM RP003187, DMRS channel estimation.  
11. Tse & Viswanath, *Fundamentals of Wireless Communication*.  
12. Project source documents: `6G_AI_CE_01.docx`, `Swetha_6G_AI_CE_Proj_01.pptx`.
"""


def write_docx(text: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    plots = ROOT / "outputs" / "plots"
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], 0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], 2)
        elif s.startswith("|"):
            doc.add_paragraph(s)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:], style="List Bullet")
        else:
            doc.add_paragraph(re.sub(r"\*\*(.+?)\*\*", r"\1", s))
    for img in ["architecture_scorecard.png", "benchmark_ls_mmse_ai.png", "cdf_nmse_estimators.png", "digital_twin_map.png", "classification_confusion_matrix.png"]:
        p = plots / img
        if p.exists():
            doc.add_picture(str(p), width=Inches(6.2))
    doc.save(path)


def write_pdf(text: str, path: Path) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    pdf.set_font("Helvetica", size=9)
    for line in text.splitlines():
        s = line.strip()
        if not s or s.startswith("```"):
            continue
        clean = s.encode("ascii", "ignore").decode("ascii")
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)[:220]
        try:
            if s.startswith("# "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.multi_cell(0, 7, clean[2:])
            elif s.startswith("## "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.multi_cell(0, 6, clean[3:])
            else:
                pdf.set_font("Helvetica", size=8)
                pdf.multi_cell(0, 4, clean)
        except Exception:
            pdf.add_page()
    pdf.output(path)


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    text = build_markdown()
    md = DOCS / "6G_AI_CHANNEL_ESTIMATION_IMPLEMENTATION.md"
    md.write_text(text, encoding="utf-8")
    write_docx(text, DOCS / "6G_AI_CHANNEL_ESTIMATION_IMPLEMENTATION.docx")
    write_pdf(text, DOCS / "6G_AI_CHANNEL_ESTIMATION_IMPLEMENTATION.pdf")
    print(f"Wrote {md}")


if __name__ == "__main__":
    main()
